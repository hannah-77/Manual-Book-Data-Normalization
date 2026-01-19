import pdfplumber
import fitz  # PyMuPDF
import os
import re
import io
import torch
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from PIL import Image
from sentence_transformers import SentenceTransformer, util
from langdetect import detect, DetectorFactory

DetectorFactory.seed = 0

# ==========================================
# 1. KONFIGURASI TRANSLASI & SKEMA
# ==========================================
TRANSLATION_MAP = {
    "1.1 Tujuan Produk/Definisi": {"ID": "1.1 Tujuan Produk/Definisi", "EN": "1.1 Intended Use/Definition"},
    "1.2 Panduan Keamanan": {"ID": "1.2 Panduan Keamanan", "EN": "1.2 Safety Guidelines"},
    "1.3 Penjelasan Simbol": {"ID": "1.3 Penjelasan Simbol", "EN": "1.3 Explanation of Symbols"},
    "1.4 Singkatan": {"ID": "1.4 Singkatan", "EN": "1.4 Abbreviations"},
    "2.0 Instalasi": {"ID": "2.0 Instalasi", "EN": "2.0 Installation"},
    "3.1 Antarmuka Pengguna": {"ID": "3.1 Antarmuka Pengguna", "EN": "3.1 User Interface"},
    "3.2 Overview": {"ID": "3.2 Overview", "EN": "3.2 Overview"},
    "3.3 Manajemen Pengguna": {"ID": "3.3 Manajemen Pengguna", "EN": "3.3 User Management"},
    "3.4 Prosedur Pemantauan": {"ID": "3.4 Prosedur Pemantauan", "EN": "3.4 Monitoring Procedure"},
    "3.5 Perhitungan Medis": {"ID": "3.5 Perhitungan Medis", "EN": "3.5 Medical Calculation"},
    "3.6 Manajemen Rekaman & Tinjauan Hasil": {"ID": "3.6 Manajemen Rekaman & Tinjauan Hasil", "EN": "3.6 Record Management & Review"},
    "4.1 Inspeksi Umum": {"ID": "4.1 Inspeksi Umum", "EN": "4.1 General Inspection"},
    "4.2 Pemeliharaan": {"ID": "4.2 Pemeliharaan", "EN": "4.2 Maintenance"},
    "4.3 Perawatan": {"ID": "4.3 Perawatan", "EN": "4.3 Care"},
    "4.4 Pembersihan": {"ID": "4.4 Pembersihan", "EN": "4.4 Cleaning"},
    "5.0 Pemecahan Masalah": {"ID": "5.0 Pemecahan Masalah", "EN": "5.0 Troubleshooting"},
    "6.1 Spesifikasi": {"ID": "6.1 Spesifikasi", "EN": "6.1 Specification"},
    "6.2 Kepatuhan Standar": {"ID": "6.2 Kepatuhan Standar", "EN": "6.2 Standard Compliance"},
    "7.1 Garansi": {"ID": "7.1 Garansi", "EN": "7.1 Warranty"},
    "7.2 Informasi Kontak": {"ID": "7.2 Informasi Kontak", "EN": "7.2 Contact Information"}
}

MY_SCHEMA_LABELS = {
    # BAB 1: PENDAHULUAN & KEAMANAN
    "1.1 Tujuan Produk/Definisi": {
        "keywords": ["tujuan produk", "definisi", "intended use", "intended purpose"],
        "exclude": []
    },
    "1.2 Panduan Keamanan": {
        "keywords": ["panduan keamanan", "safety guidelines", "peringatan", "caution", "warning", "bahaya"],
        "exclude": []
    },
    "1.3 Penjelasan Simbol": {
        "keywords": ["penjelasan simbol", "simbol", "symbols", "explanation of symbols", "arti simbol"],
        "exclude": []
    },
    "1.4 Singkatan": {
        "keywords": ["singkatan", "abbreviations", "daftar singkatan", "glossary"],
        "exclude": ["XIII", "XII", "XI", "Garansi", "Warranty"]
    },

    # BAB 2: INSTALASI
    "2.0 Instalasi": {
        "keywords": ["instalasi", "pemasangan", "installation", "setup", "persyaratan lingkungan"],
        "exclude": []
    },

    # BAB 3: PANDUAN OPERASIONAL
    "3.1 Antarmuka Pengguna": {
        "keywords": ["antarmuka pengguna", "user interface", "display", "tampilan layar", "tombol"],
        "exclude": []
    },
    "3.2 Overview": {
        "keywords": ["overview", "gambaran umum", "product overview", "accessories", "aksesoris", "perlengkapan"],
        "exclude": ["spesifikasi", "specification"]
    },
    "3.3 Manajemen Pengguna": {
        "keywords": ["manajemen pengguna", "user management", "data pasien", "pengguna"],
        "exclude": []
    },
    "3.4 Prosedur Pemantauan": {
        "keywords": ["prosedur pemantauan", "monitoring procedure", "langkah pemantauan", "cara mengukur"],
        "exclude": []
    },
    "3.5 Perhitungan Medis": {
        "keywords": ["perhitungan medis", "medical calculation", "kalkulasi dosis", "formula"],
        "exclude": []
    },
    "3.6 Manajemen Rekaman & Tinjauan Hasil": {
        "keywords": ["manajemen rekaman", "record management", "tinjauan hasil", "historical data"],
        "exclude": []
    },

    # BAB 4: PERAWATAN & PEMELIHARAAN
    "4.1 Inspeksi Umum": {
        "keywords": ["inspeksi umum", "general inspection", "pemeriksaan fisik"],
        "exclude": ["Instruksi pengukuran"]
    },
    "4.2 Pemeliharaan": {
        "keywords": ["maintenance", "pemeliharaan", "kalibrasi", "servis berkala", "penggantian suku cadang"],
        "exclude": ["penyimpanan", "storage", "cara menjaga"]
    },
    "4.3 Perawatan": {
        "keywords": ["perawatan", "care of device", "penyimpanan", "storage", "penanganan alat"],
        "exclude": ["kalibrasi", "servis teknis", "suku cadang", "maintenance"]
    },
    "4.4 Pembersihan": {
        "keywords": ["pembersihan", "cleaning", "disinfeksi", "sterilisasi"],
        "exclude": []
    },

    # BAB 5: PEMECAHAN MASALAH
    "5.0 Pemecahan Masalah": {
        "keywords": ["pemecahan masalah", "troubleshooting", "error codes", "solusi masalah"],
        "exclude": []
    },

    # BAB 6: SPESIFIKASI
    "6.1 Spesifikasi": {
        "keywords": ["spesifikasi", "specification", "technical data", "berat", "dimensi", "rentang terukur"],
        "exclude": ["accessories", "aksesoris", "perlengkapan"]
    },
    "6.2 Kepatuhan Standar": {
        "keywords": ["kepatuhan standar", "standard compliance", "IEC", "EMC", "ISO"],
        "exclude": []
    },

    # BAB 7: GARANSI & KONTAK
    "7.1 Garansi": {
        "keywords": ["garansi", "warranty", "layanan purna jual"],
        "exclude": ["IEC", "EMC", "ISO"]
    },
    "7.2 Informasi Kontak": {
        "keywords": ["informasi kontak", "contact information", "layanan pelanggan", "telepon", "alamat"],
        "exclude": []
    }
}

class VisualNormalizer:
    def __init__(self, schema):
        print("Memuat Model AI...")
        self.model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
        self.schema = schema
        self.label_keys = list(schema.keys())
        desc = [f"{k} " + " ".join(v["keywords"]) for k, v in schema.items()]
        self.embeddings = self.model.encode(desc, convert_to_tensor=True)
        self.threshold = 0.55

    def is_garbage(self, text):
        """Membuang fragmen teks vertikal yang berantakan """
        clean = text.strip()
        if len(clean) < 2: return True
        if clean.count(' ') > (len(clean) / 2): return True
        return False

    def crop_high_res(self, page_fitz, bbox, zoom=4):
        """Mencrop area simbol dan menjernihkannya (DPI Tinggi)"""
        mat = fitz.Matrix(zoom, zoom)
        pix = page_fitz.get_pixmap(matrix=mat, clip=bbox)
        return io.BytesIO(pix.tobytes("png"))

    def process_to_word(self, pdf_path, output_docx):
        lang_detected = "Indonesia" if "id" in pdf_path.name.lower() else "English"
        lang_code = "ID" if lang_detected == "Indonesia" else "EN"
        
        doc = Document()
        doc.add_heading(f'Normalisasi Visual: {pdf_path.name}', 0)
        
        pdf_fitz = fitz.open(pdf_path)
        added_headings = set()
        current_section = None

        with pdfplumber.open(pdf_path) as pdf_plumb:
            for i, page_plumb in enumerate(pdf_plumb.pages):
                print(f"  Menganalisis Halaman {i+1}...")
                page_fitz = pdf_fitz[i]
                
                # 1. Deteksi Judul via AI & Hard Match
                lines = page_plumb.extract_text().split('\n') if page_plumb.extract_text() else []
                for line in lines:
                    if self.is_garbage(line): continue
                    
                    # Logika Matching
                    matched = None
                    for k, v in self.schema.items():
                        if any(kw.lower() in line.lower() for kw in v["keywords"]):
                            matched = k; break
                    
                    if matched and matched not in added_headings:
                        current_section = matched
                        trans = TRANSLATION_MAP.get(current_section, {}).get(lang_code, current_section)
                        doc.add_heading(trans, level=1)
                        added_headings.add(current_section)
                    elif current_section:
                        doc.add_paragraph(line.strip())

                # 2. Tabel & Simbol (Cropping Visual)
                tables = page_plumb.find_tables()
                for table_obj in tables:
                    rows = table_obj.rows
                    if not rows: continue
                    
                    # PERBAIKAN: Mengambil jumlah kolom dari baris pertama
                    num_cols = len(rows[0].cells)
                    word_table = doc.add_table(rows=len(rows), cols=num_cols)
                    word_table.style = 'Table Grid'
                    word_table.autofit = True # Mengaktifkan fitur autofit agar tidak terpotong

                    for r_idx, row in enumerate(rows):
                        for c_idx, cell_bbox in enumerate(row.cells):
                            if cell_bbox is None: continue
                            
                            w_cell = word_table.cell(r_idx, c_idx)
                            
                            # Jika di bab Simbol (1.3), lakukan High-Res Crop untuk Piktogram
                            if current_section == "1.3 Penjelasan Simbol" and c_idx in [0, 2]:
                                try:
                                    img_stream = self.crop_high_res(page_fitz, cell_bbox)
                                    # Bersihkan paragraf default sebelum menambah gambar
                                    w_cell.paragraphs[0].clear()
                                    run = w_cell.paragraphs[0].add_run()
                                    run.add_picture(img_stream, width=Inches(0.40))
                                except Exception as e:
                                    print(f"      Gagal crop simbol: {e}")
                            else:
                                # Ekstrak teks untuk kolom keterangan/arti
                                cell_text = page_plumb.within_bbox(cell_bbox).extract_text()
                                w_cell.text = cell_text if cell_text else ""
                            
                            w_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc.save(output_docx)
        print(f"--- Selesai: {output_docx} ---")

if __name__ == "__main__":
    folder_in = Path(r"data_input\FOX_BABY")
    targets = ["manual_fox_baby_en.pdf", "manual_fox_baby_id.pdf"]
    recon = VisualNormalizer(MY_SCHEMA_LABELS)
    
    for target in targets:
        path = folder_in / target
        if path.exists():
            out = f"data_output/VISUAL_RECON_{path.stem}.docx"
            recon.process_to_word(path, out)